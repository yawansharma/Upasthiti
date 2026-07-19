from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def main():
    doc = Document()

    # Define custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 1. Cover Page
    doc.add_paragraph('\n\n\n')
    
    title = doc.add_paragraph('Upasthiti')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_font = title.runs[0].font
    title_font.size = Pt(36)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 102, 204) # A nice corporate blue

    slogan = doc.add_paragraph('Secure, smart, verified.')
    slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    slogan_font = slogan.runs[0].font
    slogan_font.size = Pt(16)
    slogan_font.italic = True

    doc.add_paragraph('\n\n\n\n\n\n\n\n\n')
    
    company = doc.add_paragraph('Developed by:\nNavonmesh Samadhan LLP')
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_font = company.runs[0].font
    company_font.size = Pt(14)
    company_font.bold = True

    doc.add_page_break()

    # 2. Executive Summary
    add_heading(doc, 'Executive Summary', level=1)
    doc.add_paragraph(
        "Upasthiti is an AI-powered, multi-role campus and workforce management platform designed "
        "to deliver tamper-proof attendance tracking, granular access controls, and real-time operational "
        "visibility. Built with state-of-the-art facial recognition and GPS geofencing, Upasthiti eliminates "
        "proxy attendance and ensures employees and personnel are exactly where they need to be. "
        "Designed specifically for corporate HR directors and enterprise management, the platform offers "
        "six specialized portals to handle everything from leave management to security auditing."
    )

    doc.add_paragraph()

    # 3. Core Features
    add_heading(doc, 'Core Technical Features & Capabilities', level=1)

    add_heading(doc, 'Tamper-Proof Biometric Verification', level=2)
    doc.add_paragraph(
        "• AI-Powered Face Verification: Enforces mandatory facial recognition checks powered by Hugging Face ML models before logging attendance.\n"
        "• GPS Geofencing Verification: Ensures personnel are physically within designated boundary zones (configurable from 30m to 500m) before attendance is accepted.\n"
        "• Liveness & Entry Status: Tracks if the entry was Early, Within Window, or Late, complete with timestamped selfie photo logs stored securely."
    )

    add_heading(doc, 'Comprehensive HR & Office Administration', level=2)
    doc.add_paragraph(
        "• Granular Leave Management: Fully automated leave request system with hierarchical approval chains (Level N requests go to Level N+1).\n"
        "• Real-Time Analytics & Reporting: Generate attendance and behavior reports instantly. Export data seamlessly to CSV, Excel (.xlsx), and PDF formats.\n"
        "• Employee Directory & Biometrics: Searchable employee directory with capabilities to manage approvals, suspend accounts, and securely re-enroll biometric face data."
    )

    add_heading(doc, 'Advanced Hierarchical Access Control (RBAC)', level=2)
    doc.add_paragraph(
        "Upasthiti employs a strict Role-Based Access Control system to ensure data privacy and operational security. Key roles include:\n"
        "• Employee Hub: Seamless onboarding, real-time schedule tracking, and direct leave applications.\n"
        "• L1/L2/L3 Management: Tiered supervision allowing department heads to manage their respective divisions seamlessly.\n"
        "• Security Admin: Dedicated portal for monitoring audit logs, detecting anomalies, and managing physical access control.\n"
        "• Super Admin: Top-level personnel management, organizational chart visualization, system-wide overrides, and a dedicated 'Supervision Mode' to inspect any lower-level dashboard."
    )

    add_heading(doc, 'Secure Communication & Logistics', level=2)
    doc.add_paragraph(
        "• QR-Based Distribution Events: Manage physical distributions (e.g., equipment, documents) using unique QR codes. Admins can scan recipient QR codes to mark items as distributed in real-time.\n"
        "• Internal Corporate Messaging: Secure, encrypted 1:1 Direct Messaging and department-wide broadcast channels with file attachment support (PDF, Excel, Images)."
    )

    # 4. Technical Architecture
    add_heading(doc, 'Technical Architecture & Security Options', level=1)
    
    doc.add_paragraph(
        "Upasthiti is built on a modern, highly scalable technology stack:\n"
    )
    doc.add_paragraph("• Frontend Framework: Cross-platform Flutter engine (supports Android, iOS, Windows, and Web).", style='List Bullet')
    doc.add_paragraph("• Backend Infrastructure: Powered by Appwrite Cloud NoSQL databases with real-time WebSocket syncing.", style='List Bullet')
    doc.add_paragraph("• Data Security: Employs SHA-256 cryptographic hashing for sensitive data.", style='List Bullet')
    doc.add_paragraph("• Active Account Management: Automated cleanup of dormant accounts to ensure data hygiene.", style='List Bullet')

    doc.add_page_break()

    # 5. Contact / Call to action
    add_heading(doc, 'Ready to Upgrade Your Operations?', level=1)
    doc.add_paragraph(
        "Navonmesh Samadhan LLP is committed to providing smart, secure, and verified solutions "
        "tailored to your organization's unique requirements."
    )
    doc.add_paragraph("\n\nFor technical demonstrations and integration options, please contact us:")
    doc.add_paragraph("Navonmesh Samadhan LLP\n[Contact Email Placeholder]\n[Contact Phone Placeholder]")

    # Save document
    doc.save('Upasthiti_Platform_Overview.docx')

if __name__ == '__main__':
    main()
