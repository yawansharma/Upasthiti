from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_para(doc, text, style='Normal', justify=True):
    p = doc.add_paragraph(text, style=style)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def main():
    doc = Document()

    # Define custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 1. Title Page
    doc.add_paragraph('\n\n\n\n')
    
    title = doc.add_paragraph('Comprehensive Technical & Functional Specification')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_font = title.runs[0].font
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('\nUpasthiti: AI-Powered Workforce Management')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_font = subtitle.runs[0].font
    subtitle_font.size = Pt(22)
    subtitle_font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph('\n\n\n\n\n\n\n\n')
    
    company = doc.add_paragraph('Prepared by:\nNavonmesh Samadhan LLP\n\nVersion: 1.0\nStatus: Final\nConfidentiality: Internal / Client Facing')
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_font = company.runs[0].font
    company_font.size = Pt(12)
    company_font.bold = True

    doc.add_page_break()

    # 2. Table of Contents (Placeholder)
    add_heading(doc, 'Table of Contents', level=1)
    add_para(doc, "1. Executive Summary ........................................................................................ 3")
    add_para(doc, "2. Business Case & Value Proposition ................................................................ 4")
    add_para(doc, "3. High-Level System Architecture ...................................................................... 6")
    add_para(doc, "4. Detailed Functional Modules ............................................................................ 8")
    add_para(doc, "   4.1 Employee/Student Hub .............................................................................. 8")
    add_para(doc, "   4.2 Office Administration Portal ...................................................................... 10")
    add_para(doc, "   4.3 HR Administration Portal .......................................................................... 12")
    add_para(doc, "   4.4 Security Administration Portal .................................................................. 13")
    add_para(doc, "   4.5 Logistics & Event Administration Portal .................................................. 14")
    add_para(doc, "   4.6 Super Admin / Institution Head Portal ...................................................... 15")
    add_para(doc, "5. Core Technologies Deep Dive ........................................................................ 16")
    add_para(doc, "6. Secure Communication System .................................................................... 18")
    add_para(doc, "7. Security, Privacy & Compliance .................................................................... 19")
    add_para(doc, "8. Database Schema Overview ......................................................................... 21")
    add_para(doc, "9. Deployment, Scalability & Roadmap ........................................................... 22")
    add_para(doc, "10. Conclusion & Contact .................................................................................. 23")
    
    doc.add_page_break()

    # Generate Content
    sections = [
        ("1. Executive Summary", 
         "Upasthiti is an enterprise-grade, AI-powered workforce and campus management platform engineered to deliver unprecedented operational transparency, strict access control, and robust attendance verification. Developed by Navonmesh Samadhan LLP, Upasthiti represents a paradigm shift from traditional, easily exploitable proxy attendance systems (such as legacy RFID cards or standalone fingerprint scanners) toward a dynamic, multi-factor verification matrix. By synchronously integrating state-of-the-art facial recognition machine learning models with highly precise GPS geofencing, Upasthiti ensures that personnel are authenticated not just by who they are, but exactly where they are at any given moment in time.\n\n"
         "The platform is built upon a highly scalable, cross-platform architecture utilizing the Flutter framework for client-side applications across Android, iOS, Windows, and Web. The backend infrastructure is securely hosted on Appwrite Cloud, leveraging NoSQL document databases, secure file storage buckets, and real-time WebSocket synchronization to provide instant data propagation across the entire organization. From the moment an employee submits an attendance request, the telemetry data is evaluated, verified, and distributed to relevant administrative dashboards within milliseconds.\n\n"
         "Designed specifically to cater to the complex hierarchical structures of modern corporate enterprises and large-scale educational institutions, Upasthiti features six entirely distinct, role-based administrative portals. Each portal is meticulously tailored to the specific workflows of different organizational units: Office Administration for daily approvals, HR Administration for comprehensive leave management, Security Administration for physical access auditing and anomaly detection, Logistics Administration for QR-based asset distribution, and a Super Admin dashboard for ultimate organizational oversight. This document serves as the comprehensive technical and functional specification for the Upasthiti platform, detailing every module, architectural decision, and security protocol implemented to safeguard enterprise operations.", 3),
        
        ("2. Business Case & Value Proposition",
         "The modern enterprise faces significant challenges in workforce management, primarily centering around time theft, proxy attendance, and the administrative overhead associated with manual verification and leave management. Traditional biometric systems often suffer from spoofing vulnerabilities, require expensive proprietary hardware installation at every entry point, and lack real-time synchronization with central HR systems. Furthermore, RFID and smart card systems are frequently exchanged between employees, rendering access logs entirely unreliable for serious security audits.\n\n"
         "Upasthiti systematically dismantles these challenges through a software-first approach that transforms every employee's smartphone into a secure, tamper-proof terminal. By enforcing a dual-layer verification protocol during every attendance event, the system guarantees authenticity. First, the application queries the device's GPS hardware, calculating the exact geodesic distance between the user's coordinates and the dynamically configured boundary of their designated workplace. If the user is outside the geofence, the request is immediately rejected at the client level, preventing server load and ensuring compliance.\n\n"
         "If the spatial requirement is met, the system initiates a real-time biometric challenge. The user captures a live selfie, which is transmitted securely to a dedicated Hugging Face machine learning backend. This API compares the live capture against the cryptographically secured baseline enrollment photo stored in the Appwrite Storage bucket. Only upon successful verification from both the spatial and biometric engines is the attendance log committed to the immutable database.\n\n"
         "This approach provides immense value to corporate HR directors and organizational leaders:\n"
         "• Elimination of Capital Expenditure: No proprietary biometric hardware or RFID scanners are required. The system leverages existing mobile devices.\n"
         "• Eradication of Proxy Attendance: The combination of spatial and biometric verification mathematically eliminates the possibility of buddy punching.\n"
         "• Drastic Reduction in Administrative Overhead: Automated attendance reconciliation, digitized hierarchical leave approvals, and instant report generation (CSV/Excel/PDF) free HR personnel to focus on strategic initiatives rather than manual data entry.\n"
         "• Enhanced Operational Security: Real-time anomaly detection and strict Role-Based Access Control (RBAC) ensure that sensitive organizational data remains siloed and protected against unauthorized lateral access.\n"
         "• Agile Logistics Tracking: The integrated QR distribution system digitizes the handover of physical assets (laptops, documents, access keys), providing an instant, auditable trail of custody.", 4),
        
        ("3. High-Level System Architecture",
         "The architecture of Upasthiti is designed for maximum resilience, low latency, and infinite horizontal scalability. It employs a modern microservices-inspired decoupled architecture where the client application, backend database services, and machine learning inference engines operate entirely independently, communicating via secure RESTful APIs and WebSocket protocols.\n\n"
         "3.1 Client Application (Frontend)\n"
         "The client interface is developed using Flutter (Dart 3.x), compiling to native ARM code for Android and iOS devices, and optimizing for desktop execution on Windows. This single-codebase approach guarantees absolute feature parity across all platforms while significantly reducing maintenance overhead. The UI layer implements a centralized theme engine ('AppTheme') utilizing consistent typography (Poppins), specialized micro-animations (RisingSheet transitions), and hero routing to provide a premium, fluid user experience.\n\n"
         "3.2 Backend as a Service (BaaS)\n"
         "Upasthiti relies on Appwrite Cloud for its core backend infrastructure. Appwrite provides a secure, scalable NoSQL database optimized for high-throughput read/write operations. \n"
         "• Databases: Structured document collections store users, classes, logs, messages, and distribution events.\n"
         "• Storage Buckets: Isolated storage zones manage high-resolution profile pictures, compressed attendance selfies, and encrypted community file attachments.\n"
         "• Realtime Engine: Appwrite's WebSocket implementation ('subscribe' method) is deeply integrated into the client state management, allowing administrative dashboards and community chat interfaces to update instantaneously without polling.\n\n"
         "3.3 Machine Learning Inference Engine\n"
         "Facial verification is offloaded to a specialized Python-based backend hosted on Hugging Face Spaces. This separation of concerns prevents the primary database server from being bottlenecked by computationally expensive matrix multiplications required for facial feature extraction. The ML engine exposes two highly optimized endpoints: POST /register-face and POST /login-face. It processes multipart form data, extracts facial embeddings, calculates cosine similarity against registered templates, and returns boolean verification results with sub-second latency.\n\n"
         "3.4 Geolocation Engine\n"
         "Spatial calculations are handled client-side using the `geolocator` and `latlong2` libraries. Admins define boundaries using an interactive OpenStreetMap interface (`flutter_map`). The boundary metadata (latitude, longitude, radius in meters) is serialized as JSON and stored in the class document. During attendance, the Haversine formula is applied locally to determine the precise distance between the user's current GPS fix and the boundary epicenter.", 3),
         
        ("4. Detailed Functional Modules", 
         "Upasthiti is partitioned into distinct functional modules, each safeguarded by the RBAC engine. Upon successful authentication (and CAPTCHA verification for administrators), the routing controller dynamically constructs the application topology based on the user's assigned role and hierarchical level.\n\n", 1),
         
        ("4.1 Employee / Student Hub",
         "The Employee Hub is the primary interface for standard users within the organization. It is engineered for maximum simplicity and frictionless operation.\n\n"
         "Registration & Onboarding Wizard:\n"
         "New personnel initiate the onboarding process through a multi-step registration wizard. The user inputs standard biographic data (Name, Unique ID, Department) and selects a customized security question for self-serve password recovery. Crucially, the wizard mandates the capture of a high-resolution baseline profile photo and performs the initial biometric face enrollment via the ML API. The device's GPS coordinates are also captured during registration to flag anomalous remote enrollments. Once submitted, the account enters a 'pending' state, locking all access until formal approval by an Office or HR Administrator.\n\n"
         "Dashboard & Schedule Tracking:\n"
         "Upon login, users are presented with a dynamic dashboard organizing their assigned divisions or classes. The dashboard cross-references the current system time against active 'periods' (scheduled time windows). If a session is currently active (e.g., within 10 minutes of the start or end time), the session card transforms, prompting the user to submit their attendance.\n\n"
         "Attendance Execution Flow:\n"
         "When initiating attendance, the application first requests a high-accuracy GPS fix. If the user is outside the defined geofence, a visual error is presented, and the flow terminates. If within the boundary, the front-facing camera is activated. The user captures a selfie, which is transmitted to the ML backend. Upon mathematical confirmation of identity, the image is uploaded to the Appwrite 'attendance_photos' bucket, and a permanent, immutable log is written to the database containing the timestamp, geographic status, biometric status, and calculated entry status (Early, Within Window, or Late).\n\n"
         "Self-Service Leave Application:\n"
         "Users can submit formal leave requests directly from the app. They select a leave category (e.g., Medical, Casual, Paid Leave, LTC), define the date range, and provide a justification. The request is immediately routed to their direct supervisor (Level N+1) in the administrative hierarchy, providing complete transparency into the approval status.", 4),

        ("4.2 Office Administration Portal",
         "The Office Administration module is designed for the personnel responsible for daily operational data entry, biometric management, and record keeping.\n\n"
         "Account Approval Workflow:\n"
         "Office Admins receive real-time notifications of pending registrations. They review the submitted biographic data alongside the captured profile photo to verify identity against corporate HR records. They can filter pending requests by department and execute bulk approvals or rejections, instantly transitioning user accounts from 'pending' to 'active'.\n\n"
         "Biometric Lifecycle Management:\n"
         "In cases where an employee's appearance changes significantly or the initial enrollment was flawed, the Office Admin utilizes the Biometrics tab. This interface allows the admin to bypass standard registration and forcefully re-enroll a user's facial embeddings directly into the ML backend, ensuring the system remains accurate over time.\n\n"
         "Granular Data Export & Reporting:\n"
         "The Reports tab provides advanced querying capabilities over the global attendance log collection. Admins can filter by date ranges, departments, specific employees, and attendance statuses. The module utilizes the `csv`, `excel`, and `pdf` Dart packages to compile these queries into perfectly formatted downloadable reports, ready for ingestion into legacy payroll systems or executive review.\n\n"
         "Individual Attendance Tracking:\n"
         "Office Admins have access to a specialized 'Student/Employee Directory' that allows them to drill down into the complete historical attendance record of any single individual across all their assigned divisions, providing a holistic view of punctuality and presence.", 4),

        ("4.3 HR Administration Portal",
         "The HR Administration Portal elevates oversight from daily data entry to strategic workforce management and organizational compliance.\n\n"
         "Hierarchical Leave Processing:\n"
         "HR Admins possess supreme authority over the leave management pipeline. While line managers (L2/L3 Admins) may provide initial approval, HR Admins review aggregated leave data to ensure compliance with corporate policy, track remaining leave balances, and issue final authorizations. The system maintains an immutable audit trail of who approved the leave and at what exact timestamp.\n\n"
         "Organizational Analytics:\n"
         "The HR dashboard aggregates daily attendance logs into macroscopic analytics. HR directors can instantly visualize department-wide absentee rates, identify chronic late-arrivals, and monitor overall workforce utilization through real-time generated charts and statistical summaries.\n\n"
         "Personnel Lifecycle Management:\n"
         "HR Admins collaborate with Super Admins to manage offboarding. While dormant accounts are automatically purged, HR can manually trigger account suspensions, revoke access to specific departments, and archive historical records in compliance with data retention policies.", 3),

        ("4.4 Security Administration Portal",
         "The Security Administration module focuses entirely on system integrity, access auditing, and the identification of anomalous behavior patterns.\n\n"
         "Immutable Audit Logging:\n"
         "Every critical action within Upasthiti—from admin logins and password overrides to manual attendance modifications and QR scan events—is meticulously logged. Security Admins interface with a read-only, aggressively indexed dashboard that displays this chronological audit trail, ensuring complete non-repudiation of all administrative actions.\n\n"
         "Anomaly Detection & Alerting:\n"
         "The portal continuously monitors for suspicious metadata. For example, if a user attempts to log attendance from a GPS coordinate that mathematically impossible to reach given their previous log's location and timestamp (velocity anomaly), or if an administrator experiences multiple failed CAPTCHA challenges, the Security dashboard flags these events for immediate review.\n\n"
         "Physical Access Control Integration:\n"
         "Security Admins manage the metadata linking digital identities to physical zones. By reviewing attendance density and boundary violations, they can dynamically recommend adjustments to geofence radii to optimize campus security perimeters.", 3),

        ("4.5 Logistics & Event Administration Portal",
         "The Logistics Administration module digitizes the chaotic process of distributing physical assets (laptops, onboarding kits, secure documents, or event tickets) to large groups of people.\n\n"
         "Event Lifecycle Management:\n"
         "Logistics Admins create 'Distribution Events' defining the title, venue, and scheduled datetime. The event exists in a 'draft' state while the admin configures the parameters. Recipients are added either manually via the directory search or bulk-imported by parsing standard Excel (.xlsx) files.\n\n"
         "Delegated Scanning Authority:\n"
         "To handle massive throughput (e.g., distributing 1,000 laptops in a single day), the event creator can assign multiple other administrators as authorized scanners for that specific event.\n\n"
         "Real-Time QR Verification:\n"
         "Every user generates a unique, cryptographically versioned QR code via their Employee Hub. Authorized scanning admins utilize the `mobile_scanner` library to read these codes. Upon scan, the system instantly evaluates the payload, cross-references the event's recipient list, and returns a distinct status: Success (Item Issued), Already Issued (Duplicate attempt prevented), Not Authorized, or Ineligible. Progress bars on the admin dashboard update via WebSockets, displaying the live ratio of issued items versus total recipients.", 3),

        ("4.6 Super Admin / Institution Head Portal",
         "The Super Admin (referred to as 'Dean' in academic contexts) possesses God-mode privileges over the entire Upasthiti ecosystem. This portal is concealed behind a multi-tap easter egg on the login screen to obscure its existence from unauthorized personnel.\n\n"
         "Total Personnel Control:\n"
         "The Super Admin is responsible for creating all other administrator accounts. They define the username, assign the hierarchical level (L1/L2/L3 or specialized roles), and allocate the admin to a specific department. In emergency scenarios, the Super Admin can instantly override passwords, unilaterally suspend accounts, or initiate permanent deletion protocols.\n\n"
         "Interactive Organizational Chart:\n"
         "The portal features a dynamically generated tree-view visualization of the entire administrative hierarchy. By parsing the boundary JSON metadata across all active divisions, the system calculates the complex web of head/supervisor relationships, allowing the Super Admin to visually comprehend the chain of command at a glance.\n\n"
         "Omniscient Supervision Mode:\n"
         "Unique to the Super Admin is the 'Supervision Mode'. This highly privileged function allows the Super Admin to temporarily assume the context of any subordinate administrator. By entering Supervision Mode, the Super Admin views exactly what that admin sees—their specific classes, their unread messages, their pending approvals—allowing for targeted audits and direct intervention in departmental bottlenecks.", 3),

        ("5. Core Technologies Deep Dive",
         "5.1 AI Facial Verification Engine (Hugging Face Integration)\n"
         "The biometric authentication is powered by a custom Python backend utilizing cutting-edge deep learning models for facial recognition. When a selfie is captured, the image is compressed and transmitted as multipart/form-data. The ML engine performs Face Detection (locating the bounding box of the face), Alignment (normalizing the angle), and Feature Extraction (generating a high-dimensional mathematical vector representing the unique topology of the face). This embedding is compared against the stored enrollment vector using cosine similarity. If the similarity score exceeds a strictly calibrated threshold, verification succeeds. This mathematical approach renders the system highly resistant to spoofing attempts using 2D photographs.\n\n"
         "5.2 Dynamic GPS Geofencing Engine\n"
         "The geospatial verification relies on high-accuracy GPS hardware polling. Administrators define boundaries using lat/lng coordinates and a radius in meters. During attendance, the client application utilizes the `geolocator` package to acquire the user's position, ensuring it rejects cached or low-accuracy cellular tower triangulations. The `latlong2` package executes the Haversine formula to compute the spherical distance. The system includes logic to handle edge cases, such as GPS drift, by implementing a minor tolerance buffer while strictly rejecting mock location providers.\n\n"
         "5.3 Real-Time WebSocket Synchronization\n"
         "Upasthiti leverages Appwrite's Realtime API to eliminate manual refreshing and polling. The Flutter client opens a persistent WebSocket connection upon login. Pages dynamically subscribe to specific document channels (e.g., `databases.[dbId].collections.attendance_logs.documents`). When any client writes a new log to the database, the Appwrite server immediately broadcasts an event to all subscribed clients. The Flutter `StatefulWidget` receives this payload and triggers a `setState` callback, instantly re-rendering the UI to display the new data. This architecture ensures that administrative dashboards are always perfectly synchronized with the ground truth.", 4),

        ("6. Secure Communication System",
         "Recognizing the need for secure, context-aware communication, Upasthiti features a deeply integrated messaging platform.\n\n"
         "Public Broadcast Channels:\n"
         "Every division or class automatically generates a public broadcast channel. Administrators utilize this channel to push urgent notifications, schedule changes, or general announcements. Messages sent by administrators are visually highlighted with specialized badges to assert authority and distinguish them from standard user chatter.\n\n"
         "Encrypted Direct Messaging (DMs):\n"
         "Users and administrators can initiate private 1:1 conversations. These DMs are strictly siloed; a user cannot message someone outside their assigned divisions. This prevents organizational spam and ensures communication remains relevant to operational duties.\n\n"
         "File Attachment Handling:\n"
         "The messaging system supports rich media sharing. Users can attach PDFs, Excel spreadsheets, and images using the `file_picker` library. Files are securely uploaded to the 'community_files' Appwrite storage bucket. The resulting URL is embedded in the message payload. When tapped, the client utilizes the `url_launcher` package to securely download and display the attachment.", 3),

        ("7. Security, Privacy & Compliance",
         "Enterprise security is foundational to the Upasthiti architecture, designed to protect sensitive PII and operational telemetry.\n\n"
         "Cryptographic Password Hashing:\n"
         "All user passwords undergo client-side cryptographic hashing using the SHA-256 algorithm via the Dart `crypto` package before transmission. The database stores only the resulting hash, ensuring that even in the catastrophic event of a database breach, user credentials remain unreadable. The system includes a dual-mode verification pipeline that transparently upgrades legacy plaintext passwords to secure hashes upon the next successful login.\n\n"
         "Role-Based Access Control (RBAC):\n"
         "Navigation and data fetching are strictly governed by RBAC. A user authenticated as a 'student' is programmatically prevented from accessing administrative routes. Furthermore, queries to the database are parameterized using the authenticated user's ID, preventing horizontal privilege escalation (e.g., a user attempting to query another user's attendance logs).\n\n"
         "Automated Data Hygiene & Dormant Account Purging:\n"
         "To comply with data minimization principles, Upasthiti implements an automated cleanup protocol. The `lastLogin` timestamp of every user is continuously tracked. Upon any administrator login, a background routine queries the database for accounts inactive for greater than 60 days. These dormant accounts, along with their associated profile pictures in the storage bucket, are automatically purged from the system, reducing attack surfaces and minimizing storage costs.\n\n"
         "CAPTCHA Defense:\n"
         "Administrative login portals are fortified with an automated CAPTCHA system requiring the transcription of a randomly generated 5-character alphanumeric string. This effectively neutralizes automated credential stuffing and brute-force attacks against high-privilege accounts.", 4),

        ("8. Database Schema Overview",
         "The Appwrite NoSQL database is architected for rapid reads and denormalized efficiency.\n\n"
         "• `users`: The central identity registry containing usernames, hashed passwords, roles, hierarchical levels, profile picture IDs, and security questions.\n"
         "• `classes` / divisions: Represents organizational units. Contains arrays of enrolled user IDs and a serialized JSON `boundary` object defining the geofence and L2/L3 supervisory assignments.\n"
         "• `attendance_logs`: Highly transactional collection recording every attendance event. Fields include user ID, division ID, precise timestamp, geofence boolean, facial verification boolean, and the computed entry status.\n"
         "• `periods`: Defines the temporal windows during which attendance is permitted for specific divisions.\n"
         "• `leave_requests`: Tracks the entire lifecycle of a leave application, including the requestor, date range, approval status, and the ID of the authorizing administrator.\n"
         "• `distribution_events` & `event_recipients`: Pairs a master event document with thousands of individual recipient status documents to track QR scanning progress with high granularity.", 3),

        ("9. Deployment, Scalability & Future Roadmap",
         "Upasthiti is designed to scale effortlessly from a single office of 50 employees to a multi-national enterprise of 50,000.\n\n"
         "Deployment:\n"
         "The Flutter client compiles to standard APKs/AABs for Android, IPA for iOS, and executable binaries for Windows. The web compilation allows deployment to any standard web server. The Appwrite backend is currently optimized for the Singapore cloud region, ensuring low-latency WebSocket connections across the APAC hemisphere.\n\n"
         "Roadmap & Future Enhancements:\n"
         "• Integration with Native Biometrics: Expanding the authentication matrix to include on-device fingerprint scanning as an alternative to facial recognition.\n"
         "• Automated Absence Alerting: Implementation of server-side cron jobs to automatically dispatch SMS or Email warnings to chronic absentees.\n"
         "• Migration to Appwrite Native Sessions: Transitioning from custom cryptographic hashing to Appwrite's deeply integrated session token architecture for enhanced JWT security.\n"
         "• Advanced Analytics Dashboarding: Incorporating complex data visualization libraries to render multi-dimensional performance charts directly within the HR portal.", 2),

        ("10. Conclusion & Contact",
         "Upasthiti by Navonmesh Samadhan LLP represents the pinnacle of modern workforce management. By seamlessly fusing artificial intelligence, geolocation, and real-time synchronization into a suite of powerful, role-specific applications, it provides unparalleled security and operational efficiency.\n\n"
         "For extensive technical demonstrations, source code licensing inquiries, or custom enterprise integration discussions, please contact our engineering team:\n\n"
         "Navonmesh Samadhan LLP\n"
         "Enterprise Solutions Division\n"
         "[Contact Email Placeholder]\n"
         "[Contact Phone Placeholder]\n"
         "www.navonmesh.com", 2)
    ]

    for title, content, paras in sections:
        add_heading(doc, title, level=1)
        # Split content by double newline to maintain paragraph structure
        for paragraph_text in content.split('\n\n'):
            if paragraph_text.strip():
                add_para(doc, paragraph_text)
        # Add some padding space
        doc.add_paragraph()

    # Save document
    doc.save('Upasthiti_Comprehensive_Specification.docx')

if __name__ == '__main__':
    main()
